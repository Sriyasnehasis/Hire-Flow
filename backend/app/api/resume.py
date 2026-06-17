from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
from app.core.db import get_db
from app.core.security import security_service
from app.models.user import User
from app.models.resume import Resume
from app.services.resume_parser import ResumeParsing
from app.services import ai_service
from app.utils.pdf_gen import generate_pdf_resume
from pydantic import BaseModel
import json
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter(prefix="/resumes", tags=["Resume Management"])


class ResumeResponse(BaseModel):
    id: int
    user_id: int
    original_filename: str
    ats_score: float = None
    is_current: bool
    uploaded_at: str

    class Config:
        from_attributes = True


class AutoFillResponse(BaseModel):
    message: str
    updated_fields: dict


# Resume Builder Schemas
class EducationEntry(BaseModel):
    school: str
    degree: str
    field_of_study: str
    start_date: str
    end_date: str
    grade: Optional[str] = None
    description: Optional[str] = None


class ExperienceEntry(BaseModel):
    company: str
    position: str
    start_date: str
    end_date: str
    is_current: bool = False
    description: Optional[str] = None


class ProjectEntry(BaseModel):
    title: str
    description: str
    technologies: List[str] = []
    link: Optional[str] = None


class CertificationEntry(BaseModel):
    name: str
    issuer: str
    issue_date: str
    credential_url: Optional[str] = None


class ResumeGenerateRequest(BaseModel):
    target_job: str
    template_style: str = "professional"


class CoverLetterGenerateRequest(BaseModel):
    company_name: str
    job_title: str
    job_description: Optional[str] = ""
    tone: Optional[str] = "professional"


class ResumeBuildRequest(BaseModel):
    template: str  # "modern", "professional", "creative", "minimal"
    title: str
    summary: Optional[str] = None
    skills: List[str] = []
    education: List[EducationEntry] = []
    experience: List[ExperienceEntry] = []
    projects: List[ProjectEntry] = []
    certifications: List[CertificationEntry] = []


class ResumeUpdateRequest(BaseModel):
    title: Optional[str] = None
    summary: Optional[str] = None
    skills: Optional[List[str]] = None
    education: Optional[List[Dict[str, Any]]] = None
    experience: Optional[List[Dict[str, Any]]] = None
    projects: Optional[List[Dict[str, Any]]] = None
    certifications: Optional[List[Dict[str, Any]]] = None


@router.post("/upload")
async def upload_resume(
    file: UploadFile = File(...),
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """
    Upload and parse resume (PDF or DOCX)
    Returns parsed data and saves to database
    """
    # Validate file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename required")
    
    allowed_extensions = {'.pdf', '.docx', '.doc'}
    file_ext = '.' + file.filename.split('.')[-1].lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"File type not supported. Allowed: {allowed_extensions}"
        )
    
    try:
        # Read file content
        content = await file.read()
        
        # Parse based on file type
        if file_ext == '.pdf':
            raw_text = await ResumeParsing.parse_pdf(content)
        elif file_ext in ['.docx', '.doc']:
            raw_text = await ResumeParsing.parse_docx(content)
        
        # Extract structured data
        parsed_data = ResumeParsing.extract_structured_data(raw_text)
        
        # Mark previous resumes as not current
        db.query(Resume).filter(
            Resume.user_id == current_user_id,
            Resume.is_current == True
        ).update({"is_current": False})
        
        # Create new resume record
        resume = Resume(
            user_id=current_user_id,
            original_filename=file.filename,
            raw_text=raw_text,
            parsed_education=json.dumps(parsed_data.get("education", [])),
            parsed_experience=json.dumps(parsed_data.get("experience", [])),
            parsed_skills=json.dumps(parsed_data.get("skills", [])),
            is_current=True,
            version_number=1
        )
        
        db.add(resume)
        db.commit()
        db.refresh(resume)
        
        return {
            "message": "Resume uploaded and parsed successfully",
            "resume_id": resume.id,
            "filename": file.filename,
            "parsed_data": parsed_data,
            "raw_text_length": len(raw_text)
        }
    
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error processing resume: {str(e)}")


@router.post("/auto-fill-profile")
async def auto_fill_profile(
    resume_id: int,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """
    Auto-fill user profile fields from parsed resume data
    """
    # Get resume
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user_id
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Get user
    user = db.query(User).filter(User.id == current_user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    try:
        # Parse JSON fields
        skills = json.loads(resume.parsed_skills or '[]')
        experience = json.loads(resume.parsed_experience or '[]')
        education = json.loads(resume.parsed_education or '[]')
        
        # Update user profile
        updated_fields = {}
        
        # Auto-fill skills
        if skills:
            user.primary_skills = skills
            updated_fields["primary_skills"] = skills
        
        # Auto-fill education (take first degree if available)
        if education:
            user.educational_qualification = education[0]
            updated_fields["educational_qualification"] = education[0]
        
        # Update resume text
        user.resume_text = resume.raw_text
        
        db.commit()
        db.refresh(user)
        
        return AutoFillResponse(
            message="Profile auto-filled from resume",
            updated_fields=updated_fields
        )
    
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error auto-filling profile: {str(e)}")


@router.get("")
async def list_resumes(
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
) -> List[ResumeResponse]:
    """Get all resumes for current user"""
    resumes = db.query(Resume).filter(Resume.user_id == current_user_id).order_by(Resume.uploaded_at.desc()).all()
    return resumes


@router.get("/{resume_id}")
async def get_resume(
    resume_id: int,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """Get detailed resume information and parsed data"""
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user_id
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    return {
        "id": resume.id,
        "title": resume.title or resume.original_filename,
        "filename": resume.original_filename,
        "uploaded_at": resume.uploaded_at,
        "is_current": resume.is_current,
        "summary": resume.summary,
        "parsed_data": {
            "skills": json.loads(resume.parsed_skills or '[]'),
            "education": json.loads(resume.parsed_education or '[]'),
            "experience": json.loads(resume.parsed_experience or '[]'),
            "projects": json.loads(resume.parsed_projects or '[]'),
            "certifications": json.loads(resume.parsed_certifications or '[]')
        },
        "ats_score": resume.ats_score,
        "template": resume.template_id,
        "raw_text_preview": resume.raw_text[:500] if resume.raw_text else None
    }


# ============================================================
# RESUME BUILDER ENDPOINTS (Feature 1)
# ============================================================

@router.get("/templates")
async def list_templates():
    """List available resume templates"""
    templates = [
        {
            "id": "modern",
            "name": "Modern",
            "description": "Clean, contemporary design with bold sections",
            "color_scheme": "blue",
            "preview_url": "/templates/modern.png"
        },
        {
            "id": "professional",
            "name": "Professional",
            "description": "Classic format preferred by traditional industries",
            "color_scheme": "gray",
            "preview_url": "/templates/professional.png"
        },
        {
            "id": "creative",
            "name": "Creative",
            "description": "Standout design for creative professionals",
            "color_scheme": "purple",
            "preview_url": "/templates/creative.png"
        },
        {
            "id": "minimal",
            "name": "Minimal",
            "description": "Simple, elegant design focusing on content",
            "color_scheme": "black",
            "preview_url": "/templates/minimal.png"
        }
    ]
    return {"templates": templates}


@router.post("/create")
async def create_resume(
    request: ResumeBuildRequest,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """Create new resume from builder"""
    try:
        # Mark previous resumes as not current
        db.query(Resume).filter(
            Resume.user_id == current_user_id,
            Resume.is_current == True
        ).update({"is_current": False})
        
        # Create resume record
        resume = Resume(
            user_id=current_user_id,
            original_filename=f"{request.title}_Resume",
            created_from="built",
            template_id=request.template,
            title=request.title,
            summary=request.summary,
            parsed_skills=json.dumps(request.skills),
            parsed_education=json.dumps([e.dict() for e in request.education]),
            parsed_experience=json.dumps([e.dict() for e in request.experience]),
            parsed_projects=json.dumps([p.dict() for p in request.projects]),
            parsed_certifications=json.dumps([c.dict() for c in request.certifications]),
            is_current=True,
            version_number=1
        )
        
        db.add(resume)
        db.commit()
        db.refresh(resume)
        
        return {
            "status": "success",
            "resume_id": resume.id,
            "message": f"Resume '{request.title}' created successfully",
            "template": request.template
        }
    
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error creating resume: {str(e)}")


@router.put("/{resume_id}")
async def update_resume(
    resume_id: int,
    request: ResumeUpdateRequest,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """Update resume sections"""
    # Get resume
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user_id
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    try:
        # Update fields if provided
        if request.title:
            resume.title = request.title
        if request.summary:
            resume.summary = request.summary
        if request.skills is not None:
            resume.parsed_skills = json.dumps(request.skills)
        if request.education is not None:
            resume.parsed_education = json.dumps(request.education)
        if request.experience is not None:
            resume.parsed_experience = json.dumps(request.experience)
        if request.projects is not None:
            resume.parsed_projects = json.dumps(request.projects)
        if request.certifications is not None:
            resume.parsed_certifications = json.dumps(request.certifications)
        
        resume.updated_at = datetime.utcnow()
        resume.version_number += 1
        
        db.commit()
        db.refresh(resume)
        
        return {
            "status": "success",
            "resume_id": resume.id,
            "message": "Resume updated successfully",
            "version": resume.version_number
        }
    
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error updating resume: {str(e)}")


@router.get("/{resume_id}/preview")
async def get_resume_preview(
    resume_id: int,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """Get resume preview with all sections"""
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user_id
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    return {
        "id": resume.id,
        "title": resume.title,
        "template": resume.template_id,
        "created_from": resume.created_from,
        "summary": resume.summary,
        "skills": json.loads(resume.parsed_skills or '[]'),
        "education": json.loads(resume.parsed_education or '[]'),
        "experience": json.loads(resume.parsed_experience or '[]'),
        "projects": json.loads(resume.parsed_projects or '[]'),
        "certifications": json.loads(resume.parsed_certifications or '[]'),
        "ats_score": resume.ats_score,
        "version": resume.version_number
    }


@router.delete("/{resume_id}")
async def delete_resume(
    resume_id: int,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """Delete a resume"""
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user_id
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    try:
        db.delete(resume)
        db.commit()
        return {"status": "success", "message": "Resume deleted successfully"}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error deleting resume: {str(e)}")


@router.get("/{resume_id}/ats-analysis")
async def get_ats_analysis(
    resume_id: int,
    job_title: Optional[str] = None,
    job_description: Optional[str] = None,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """Get ATS analysis for a resume using Gemini AI"""
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user_id
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    try:
        # Get resume text
        resume_text = resume.raw_text or ""
        if not resume_text:
            # Reconstruct from structured data
            parts = []
            parts.append(resume.title or "")
            parts.append(resume.summary or "")
            parts.extend(json.loads(resume.parsed_skills or '[]'))
            for exp in json.loads(resume.parsed_experience or '[]'):
                if isinstance(exp, dict):
                    parts.append(exp.get('company', '') + ' ' + exp.get('position', ''))
            resume_text = " ".join(parts)
        
        # Analyze resume using Gemini AI Service
        analysis = await ai_service.get_ats_score(resume_text, job_description or job_title or "")
        
        # Update ATS score in database
        resume.ats_score = analysis.get("score", 0)
        resume.ats_last_checked = datetime.utcnow()
        db.commit()
        
        return {
            "resume_id": resume.id,
            "ats_score": analysis.get("score", 0),
            "feedback": analysis.get("feedback", ""),
            "matched_skills": analysis.get("matched_skills", []),
            "missing_skills": analysis.get("missing_skills", []),
            "chart_data": analysis.get("chart_data", {}),
            "last_checked": resume.ats_last_checked
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analyzing resume: {str(e)}")


@router.post("/generate-content")
async def generate_resume_ai(
    request: ResumeGenerateRequest,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """Generate resume content using Gemini AI based on user profile and target job"""
    user = db.query(User).filter(User.id == current_user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    experience = []  # You might want to parse this from user's current resume or profile
    education = []
    
    if user.resume_text:
        # Simplistic extraction for context if they have an uploaded resume
        pass

    # Call Gemini
    result = await ai_service.generate_resume_content(
        full_name=user.full_name or "Applicant",
        profession=user.profession or "",
        experience=experience,
        skills=user.primary_skills or [],
        education=education,
        target_job=request.target_job,
        template_style=request.template_style
    )
    
    if "error" in result:
        raise HTTPException(status_code=500, detail=result["error"])
        
    return result


@router.post("/generate-cover-letter")
async def generate_cover_letter_ai(
    request: CoverLetterGenerateRequest,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """Generate a cover letter using Gemini AI"""
    user = db.query(User).filter(User.id == current_user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
        
    result = await ai_service.generate_cover_letter(
        full_name=user.full_name or "Applicant",
        profession=user.profession or "",
        skills=user.primary_skills or [],
        company_name=request.company_name,
        job_title=request.job_title,
        job_description=request.job_description,
        tone=request.tone
    )
    
    return {"cover_letter": result}


# ============================================================
# EXPORT ENDPOINTS (Feature 4)
# ============================================================

def generate_docx_resume(resume: Resume) -> bytes:
    """Generate DOCX resume from resume data"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(resume.title or "Resume")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    if resume.summary:
        summary_heading = doc.add_paragraph()
        summary_heading_run = summary_heading.add_run("PROFESSIONAL SUMMARY")
        summary_heading_run.font.size = Pt(12)
        summary_heading_run.font.bold = True
        summary_heading.paragraph_format.space_before = Pt(6)
        summary_heading.paragraph_format.space_after = Pt(3)
        
        summary_text = doc.add_paragraph(resume.summary)
        summary_text.paragraph_format.space_after = Pt(9)
    
    # Skills
    skills = json.loads(resume.parsed_skills or '[]')
    if skills:
        skills_heading = doc.add_paragraph()
        skills_heading_run = skills_heading.add_run("SKILLS")
        skills_heading_run.font.size = Pt(12)
        skills_heading_run.font.bold = True
        skills_heading.paragraph_format.space_before = Pt(6)
        skills_heading.paragraph_format.space_after = Pt(3)
        
        skills_text = doc.add_paragraph(", ".join(skills))
        skills_text.paragraph_format.space_after = Pt(9)
    
    # Experience
    experience = json.loads(resume.parsed_experience or '[]')
    if experience:
        exp_heading = doc.add_paragraph()
        exp_heading_run = exp_heading.add_run("EXPERIENCE")
        exp_heading_run.font.size = Pt(12)
        exp_heading_run.font.bold = True
        exp_heading.paragraph_format.space_before = Pt(6)
        exp_heading.paragraph_format.space_after = Pt(3)
        
        for exp in experience:
            # Position and Company
            pos_para = doc.add_paragraph()
            pos_run = pos_para.add_run(exp.get('position', 'Position'))
            pos_run.font.bold = True
            pos_run.font.size = Pt(11)
            pos_para.add_run(f" at {exp.get('company', 'Company')}")
            
            # Dates
            if exp.get('start_date') or exp.get('end_date'):
                date_text = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}"
                date_para = doc.add_paragraph(date_text)
                date_para.paragraph_format.left_indent = Inches(0.25)
            
            # Description
            if exp.get('description'):
                desc_para = doc.add_paragraph(exp.get('description'))
                desc_para.paragraph_format.left_indent = Inches(0.25)
            
            doc.add_paragraph()  # Spacing
    
    # Education
    education = json.loads(resume.parsed_education or '[]')
    if education:
        edu_heading = doc.add_paragraph()
        edu_heading_run = edu_heading.add_run("EDUCATION")
        edu_heading_run.font.size = Pt(12)
        edu_heading_run.font.bold = True
        edu_heading.paragraph_format.space_before = Pt(6)
        edu_heading.paragraph_format.space_after = Pt(3)
        
        for edu in education:
            # Degree and School
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(edu.get('degree', 'Degree'))
            degree_run.font.bold = True
            degree_para.add_run(f" from {edu.get('school', 'School')}")
            
            # Field of Study
            if edu.get('field_of_study'):
                field_para = doc.add_paragraph(f"Field of Study: {edu.get('field_of_study')}")
                field_para.paragraph_format.left_indent = Inches(0.25)
            
            doc.add_paragraph()  # Spacing
    
    # Convert to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


@router.get("/{resume_id}/export/{format_type}")
async def export_resume(
    resume_id: int,
    format_type: str,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """Export resume as PDF or DOCX"""
    # Validate format
    if format_type.lower() not in ['pdf', 'docx']:
        raise HTTPException(
            status_code=400,
            detail="Format must be 'pdf' or 'docx'"
        )
    
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user_id
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    try:
        filename = f"{resume.title or 'Resume'}.{format_type.lower()}"
        
        if format_type.lower() == 'docx':
            # Generate DOCX
            docx_content = generate_docx_resume(resume)
            return FileResponse(
                io.BytesIO(docx_content),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=filename
            )
        
        elif format_type.lower() == 'pdf':
            # Generate PDF
            pdf_content = generate_pdf_resume(resume)
            return FileResponse(
                io.BytesIO(pdf_content),
                media_type="application/pdf",
                filename=filename
            )
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error exporting resume: {str(e)}"
        )


class ResumeSynthesizeRequest(BaseModel):
    target_job: str
    template: str = "professional"  # modern, professional, creative, minimal
    enrich_github: bool = True


@router.post("/synthesize-ai")
async def synthesize_resume_ai(
    request: ResumeSynthesizeRequest,
    current_user_id: int = Depends(security_service.get_user_from_token),
    db: Session = Depends(get_db)
):
    """
    Synthesizes a complete resume using Gemini AI:
    1. Fetches user details (LinkedIn bio, name, profession, qualifications)
    2. Optional: Fetches top GitHub repositories & languages
    3. Calls Gemini to generate ATS-optimized summary, skills, experience bullets, and projects list
    4. Creates and stores a new Resume node in the database
    """
    from app.services.github_service import GitHubDataService
    
    user = db.query(User).filter(User.id == current_user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
        
    # Prepare educational & experience details from profile
    education = []
    if user.educational_qualification:
        education.append({
            "school": "University Node",
            "degree": user.educational_qualification,
            "field_of_study": "Software Engineering/Science",
            "start_date": "2022",
            "end_date": "2026"
        })
        
    experience = []
    if user.current_company or user.profession:
        experience.append({
            "company": user.current_company or "Enterprise Network",
            "position": user.profession or "Tech Professional",
            "start_date": "2025",
            "end_date": "Present",
            "is_current": True,
            "description": user.bio or "Responsible for developing systems."
        })

    # Pull GitHub repositories if requested and connected
    github_projects = []
    if request.enrich_github and user.github_username and user.github_access_token:
        try:
            github_projects = await GitHubDataService.get_github_projects(user)
        except Exception as e:
            print(f"Error fetching GitHub projects for synthesis: {e}")

    # Combine into a prompt for Gemini content generation
    prompt = f"""You are a professional ATS resume writer. Generate an optimized resume dataset for:
Full Name: {user.full_name or "Professional"}
Target Job: {request.target_job}
LinkedIn Bio: {user.bio or ""}
Experience: {user.profession or ""} at {user.current_company or "None"}
Skills: {', '.join(user.primary_skills or [])}

GITHUB PROJECTS FOR INCORPORATION (if any):
{json.dumps(github_projects[:3], indent=2)}

Return a JSON object with these fields:
{{
  "professional_summary": "<compelling 3-sentence summary>",
  "skills": ["skill1", "skill2", "skill3", ...],
  "experience": [
    {{
      "company": "...",
      "position": "...",
      "start_date": "...",
      "end_date": "...",
      "description": "..."
    }}
  ],
  "projects": [
    {{
      "title": "<use repo name or project name>",
      "description": "<detailed bullet points explaining implementation and technologies>",
      "technologies": ["tech1", "tech2"],
      "link": "<repo url>"
    }}
  ]
}}

Rules:
- Formulate projects based on the GitHub data provided. Make them sound highly technical, emphasizing achievements.
- Clean up and optimize the skills list.
- Return ONLY valid JSON, no markdown."""

    try:
        # Generate content using Gemini fallback service
        result = await ai_service._call_gemini(prompt, expect_json=True)
        if not result or "professional_summary" not in result:
            raise HTTPException(status_code=500, detail="Gemini failed to output valid resume structure.")
        
        # Mark previous resumes as not current
        db.query(Resume).filter(
            Resume.user_id == current_user_id,
            Resume.is_current == True
        ).update({"is_current": False})

        # Save to database
        new_resume = Resume(
            user_id=current_user_id,
            original_filename=f"AI_Synthesized_{request.target_job.replace(' ', '_')}_Resume",
            created_from="built",
            template_id=request.template,
            title=f"{request.target_job} Resume",
            summary=result.get("professional_summary", ""),
            parsed_skills=json.dumps(result.get("skills", [])),
            parsed_education=json.dumps(education),
            parsed_experience=json.dumps(result.get("experience", [])),
            parsed_projects=json.dumps(result.get("projects", [])),
            parsed_certifications=json.dumps([]),
            is_current=True,
            version_number=1,
            file_path="local_db"
        )
        
        db.add(new_resume)
        db.commit()
        db.refresh(new_resume)
        
        # Calculate initial ATS score
        ats_analysis = await ai_service.get_ats_score(
            resume_text=f"{new_resume.title} {new_resume.summary} " + " ".join(result.get("skills", [])),
            job_description=request.target_job
        )
        new_resume.ats_score = ats_analysis.get("score", 70)
        db.commit()

        return {
            "status": "success",
            "resume_id": new_resume.id,
            "message": "Resume synthesized and stored successfully!",
            "template": request.template,
            "ats_score": new_resume.ats_score
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Resume synthesis failed: {str(e)}")
